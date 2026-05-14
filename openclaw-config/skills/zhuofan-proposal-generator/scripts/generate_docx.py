"""
ZX-01 方案初稿 → Word 文档生成器（Skill 内嵌版）

用法：
  python generate_docx.py --json proposal.json
  python generate_docx.py --json proposal.json --output /path/to/output.docx

输入 JSON 结构见 SCHEMA 常量。
依赖：pip install python-docx
"""

import json
import sys
import os
import argparse
import re
import time
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx 未安装，请执行 pip install python-docx", file=sys.stderr)
    sys.exit(1)


SCHEMA = """
输入 JSON 结构：
{
  "document_type": "solution/achievement_report",
  "project_name": "项目名称",
  "customer_name": "客户单位名称",
  "customer_type": "数据局/政务中心/大数据中心/数据集团/企业",
  "region": "区域",
  "sections": {
    "policy_background": "政策背景正文（可多段，用\\n分隔）",
    "current_status": "现状概述正文",
    "pain_points": ["痛点1", "痛点2", "痛点3"],
    "overall_goal": "总体目标正文",
    "sub_goals": ["分项目标1", "分项目标2"],
    "modules": [
      {
        "name": "模块名称",
        "content": "模块详细内容",
        "product": "匹配的卓繁产品（无则写待确认）"
      }
    ],
    "architecture": "技术架构描述（空字符串则用默认五层架构）",
    "tech_selection": "技术选型说明",
    "phases": [
      {"name": "一期", "duration": "6个月", "content": "建设内容", "deliverables": "交付物"}
    ],
    "budget_items": [
      {"name": "模块名", "amount": "待定", "note": ""}
    ],
    "highlights": ["亮点1", "亮点2"],
    "benefits": ["效益1", "效益2"],
    "cases": [
      {"name": "案例名", "source": "来源", "url": "来源链接", "status": "已核验/待核验原文", "summary": "摘要"}
    ],
    "review_notes": ["需人工审核项1", "需人工审核项2"],
    "ref_proposals": "参考的历史方案",
    "ref_products": "匹配的产品模块",
    "policy_search_status": "成功/部分成功/失败降级",
    "policy_search_note": "政策检索说明",
    "policy_sources": [
      {"name": "政策名称", "issuer": "发文单位", "date": "发文时间", "url": "来源链接", "status": "已核验/待核验原文"}
    ],
    "industry_sources": [
      {"name": "行业资料名称", "issuer": "发布单位", "date": "发布时间", "url": "来源链接", "status": "已核验/待核验原文"}
    ],
    "ref_policies_count": 0,
    "ref_cases_count": 0,
    "source_note": "政策来源、行业资料、案例来源详见来源清单；其余内容为模型辅助生成，需人工审核确认。",
    "report_overview": "建设总体概况（建设情况汇报使用）",
    "achievement_scenes": [
      {
        "name": "场景名称",
        "subtitle": "一句话亮点标题",
        "overview": "场景概述",
        "measures": ["一是具体做法与成效", "二是具体做法与成效", "三是具体做法与成效"]
      }
    ],
    "next_steps": ["下一步计划1", "下一步计划2"],
    "conclusion": "结尾总结"
  }
}
"""


FONT_SONGTI = "宋体"
FONT_FANGSONG = "仿宋_GB2312"
FONT_HEITI = "黑体"
FONT_KAITI = "楷体_GB2312"
FONT_XIAOBIAOSONG = "方正小标宋_GBK"
CN_SUBHEADINGS = ["（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）"]


# ── 工具函数 ──────────────────────────────────────────────

def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False, color: RGBColor = None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _setup_official_page(doc):
    """参考 official-doc-writer / GB/T 9704-2012 设置 A4 版面。"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def _set_paragraph_format(p, first_line_indent: bool = True, space_before: int = 0, space_after: int = 0):
    """参考公文正文格式：3号字、首行缩进2字符、1.5倍行距。"""
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(32)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)


def _set_cell_shading(cell, color_hex: str):
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(
        qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"}
    )
    tc_pr.append(shading)


def _heading(doc, text, level):
    p = doc.add_paragraph()
    _set_paragraph_format(p, first_line_indent=True, space_before=8 if level == 1 else 0)
    if level == 1:
        run = p.add_run(text)
        _set_run_font(run, FONT_HEITI, 16)
    elif level == 2:
        run = p.add_run(text)
        _set_run_font(run, FONT_KAITI, 16)
    else:
        run = p.add_run(text)
        _set_run_font(run, FONT_FANGSONG, 16)
    return p


def _strip_inline_sources(text: str) -> str:
    """正文不内嵌来源链接，来源统一放入附录来源清单。"""
    if not text:
        return ""
    text = re.sub(r"\s*[\[【]来源[:：][^\]】]+[\]】]", "", text)
    text = re.sub(r"\s*[（(]来源[:：][^）)]+[）)]", "", text)
    text = re.sub(r"\s*https?://\S+", "", text)
    return text.strip()


def _add_paragraphs(doc, text: str):
    """将含 \\n 的长文本拆成多段落写入"""
    for line in text.split("\n"):
        line = _strip_inline_sources(line.strip())
        if line:
            p = doc.add_paragraph()
            _set_paragraph_format(p)
            run = p.add_run(line)
            _set_run_font(run, FONT_FANGSONG, 16)


def _make_header_row(table, headers, bg="FFFFFF"):
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        _set_cell_shading(cell, bg)
        for p in cell.paragraphs:
            for r in p.runs:
                _set_run_font(r, FONT_HEITI, 16, True)


def _format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _set_paragraph_format(p, first_line_indent=False)
                for r in p.runs:
                    _set_run_font(r, FONT_FANGSONG, 16, bool(r.bold), r.font.color.rgb)


def _subheading(index: int, text: str) -> str:
    prefix = CN_SUBHEADINGS[index - 1] if index <= len(CN_SUBHEADINGS) else f"（{index}）"
    return f"{prefix}{text}"


def _enforce_policy_traceability(sections: dict):
    """没有政策来源链接时，自动把政策引用降级为待联网核验。"""
    policy_background = sections.get("policy_background", "")
    policy_sources = sections.get("policy_sources") or []
    ref_policies_count = sections.get("ref_policies_count", 0) or 0
    ref_count_match = re.search(r"\d+", str(ref_policies_count))
    ref_count = int(ref_count_match.group()) if ref_count_match else 0
    has_policy_claims = "《" in policy_background or ref_count > 0

    if has_policy_claims and not policy_sources:
        sections["policy_search_status"] = "失败降级"
        sections["policy_search_note"] = (
            "正文包含政策引用，但生成数据未提供可核验来源链接，已按待联网核验处理。"
        )
        sections["ref_policies_count"] = 0
        warning = "政策来源链接未随生成数据提供，以上政策引用需进一步联网核验。[待联网核验]"
        if "待联网核验" not in policy_background:
            sections["policy_background"] = f"{policy_background}\n\n{warning}".strip()


def _enforce_case_traceability(sections: dict):
    """成功案例没有来源链接时，自动标注为待联网核验。"""
    cases = sections.get("cases") or []
    for case in cases:
        if not isinstance(case, dict):
            continue
        url = (case.get("url") or "").strip()
        source = (case.get("source") or "").strip()
        has_link = url.startswith(("http://", "https://")) or "http://" in source or "https://" in source
        if has_link:
            continue
        case["status"] = case.get("status") or "待核验原文"
        case["url"] = "来源链接未提供，需联网核验。[待联网核验]"
        case["source"] = source or "公开信息/行业案例"


def _source_heading(item: dict) -> str:
    return item.get("name") or item.get("title") or item.get("source") or "未命名来源"


def _source_fields(item: dict, issuer_label: str = "发布单位"):
    return [
        (issuer_label, item.get("issuer") or item.get("source") or ""),
        ("时间", item.get("date", "")),
        ("状态", item.get("status", "")),
        ("来源链接", item.get("url", "")),
    ]


def _add_source_list(doc, title: str, items: list, issuer_label: str = "发布单位"):
    if not items:
        return
    _heading(doc, title, 2)
    for i, item in enumerate(items, 1):
        if not isinstance(item, dict):
            continue
        p = doc.add_paragraph()
        _set_paragraph_format(p, first_line_indent=False)
        run = p.add_run(f"{i}. {_source_heading(item)}")
        _set_run_font(run, FONT_HEITI, 16, True)

        for label, value in _source_fields(item, issuer_label):
            if not value:
                continue
            p = doc.add_paragraph()
            _set_paragraph_format(p)
            label_run = p.add_run(f"{label}：")
            _set_run_font(label_run, FONT_HEITI, 16, True)
            value_run = p.add_run(value)
            _set_run_font(value_run, FONT_FANGSONG, 16)


def _is_achievement_report(data: dict, sections: dict) -> bool:
    doc_type = (data.get("document_type") or sections.get("document_type") or "").lower()
    project_name = data.get("project_name", "")
    return (
        doc_type in {"achievement_report", "report", "work_report", "建设情况汇报", "成效汇报"}
        or "建设情况" in project_name
        or "汇报" in project_name
    )


def _add_common_appendices(doc, customer_type: str, region: str, sections: dict):
    doc.add_page_break()
    _heading(doc, "附录A：人工审核提示", 1)
    default_reviews = [
        "政策引用是否准确（请核对原文及文号）",
        "产品模块匹配是否与公司最新产品清单一致",
        "预算金额需由商务部门填写",
        "方案创新点建议资深顾问补充差异化设计",
        "客户承诺与量化指标需确认可达性",
        "竞对策略如需补充请人工添加",
    ]
    for item in sections.get("review_notes", default_reviews):
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        run = p.add_run(f"[ ] {item}")
        _set_run_font(run, FONT_FANGSONG, 16)

    _heading(doc, "附录B：生成说明", 1)
    info_t = doc.add_table(rows=9, cols=2)
    _format_table(info_t)
    info_data = [
        ("客户类型", customer_type),
        ("区域", region),
        ("参考历史方案", sections.get("ref_proposals", "无（知识库未接入）")),
        ("匹配产品模块", sections.get("ref_products", "无")),
        ("政策检索状态", sections.get("policy_search_status", "未说明")),
        ("政策检索说明", sections.get("policy_search_note", "未说明")),
        ("引用政策数量", f"{sections.get('ref_policies_count', 0)} 条"),
        ("引用案例数量", f"{sections.get('ref_cases_count', 0)} 个"),
        ("内容来源说明", sections.get("source_note", "政策来源、行业资料、案例来源详见来源清单；其余内容为模型辅助生成，需人工审核确认。")),
    ]
    for i, (k, v) in enumerate(info_data):
        info_t.rows[i].cells[0].text = k
        info_t.rows[i].cells[1].text = v
        for p in info_t.rows[i].cells[0].paragraphs:
            for r in p.runs:
                _set_run_font(r, FONT_HEITI, 16, True)
    _format_table(info_t)

    source_groups = [
        ("政策来源", sections.get("policy_sources", []), "发文单位"),
        ("行业资料来源", sections.get("industry_sources", []), "发布单位"),
        ("案例来源", sections.get("cases", []), "来源"),
    ]
    if any(items for _, items, _ in source_groups):
        _heading(doc, "附录C：来源清单", 1)
        for title, items, issuer_label in source_groups:
            _add_source_list(doc, title, items, issuer_label)

    disclaimer = doc.add_paragraph()
    _set_paragraph_format(disclaimer, space_before=12)
    run = disclaimer.add_run(
        "本方案初稿由「咨询顾问-卓智」数字员工辅助生成，所有内容需经人工审核确认后方可使用。"
    )
    _set_run_font(run, FONT_FANGSONG, 16)


def _save_doc(doc, output_path: str, region: str, project_name: str) -> str:
    if not output_path:
        today = date.today().strftime("%Y%m%d")
        safe = lambda t: t.replace("/", "_").replace("\\", "_").replace(" ", "")
        output_path = f"ZX01_{safe(region)}_{safe(project_name[:10])}_解决方案_初稿_{today}.docx"

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)


# ── 主函数 ────────────────────────────────────────────────

def build_proposal_docx(data: dict, output_path: str = None) -> str:
    project_name = data["project_name"]
    customer_name = data["customer_name"]
    customer_type = data["customer_type"]
    region = data["region"]
    s = data.get("sections", {})
    is_report = _is_achievement_report(data, s)
    _enforce_policy_traceability(s)
    _enforce_case_traceability(s)

    doc = Document()
    _setup_official_page(doc)

    # 全局字体
    style = doc.styles["Normal"]
    style.font.name = FONT_FANGSONG
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FANGSONG)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ───── 封面 ─────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(title_p, first_line_indent=False, space_after=8)
    if is_report:
        title_text = project_name if "汇报" in project_name else f"关于{project_name}建设情况的汇报"
    else:
        title_text = f"《{project_name}》\n解决方案"
    run = title_p.add_run(title_text)
    _set_run_font(run, FONT_XIAOBIAOSONG, 22, True)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(sub_p, first_line_indent=False, space_after=24)
    run = sub_p.add_run("（初稿）")
    _set_run_font(run, FONT_FANGSONG, 16)

    # 元信息
    meta_table = doc.add_table(rows=4, cols=4)
    _format_table(meta_table)
    meta_rows = [
        ("编制单位", "卓繁信息集团股份有限公司", "目标客户", customer_name),
        ("客户类型", customer_type, "区域", region),
        ("编制日期", date.today().strftime("%Y年%m月%d日"), "版本", "V0.1"),
        ("状态", "AI辅助生成初稿 - 待人工审核", "", ""),
    ]
    for i, (k1, v1, k2, v2) in enumerate(meta_rows):
        cells = meta_table.rows[i].cells
        cells[0].text = k1
        cells[1].text = v1
        if k2:
            cells[2].text = k2
            cells[3].text = v2
        for j in [0, 2]:
            if cells[j].text:
                for p in cells[j].paragraphs:
                    for r in p.runs:
                        _set_run_font(r, FONT_HEITI, 16, True)
    _format_table(meta_table)

    doc.add_paragraph("")

    if is_report:
        _heading(doc, "一、建设总体概况", 1)
        _add_paragraphs(doc, s.get("report_overview") or s.get("current_status", "【待补充建设总体概况】"))

        _heading(doc, "二、场景建设成效", 1)
        scenes = s.get("achievement_scenes") or []
        if scenes:
            ordinals = ["一是", "二是", "三是", "四是", "五是", "六是", "七是"]
            for i, scene in enumerate(scenes, 1):
                title = scene.get("name", f"场景{i}")
                subtitle = scene.get("subtitle", "")
                _heading(doc, f"{title}：{subtitle}" if subtitle else title, 2)
                if scene.get("overview"):
                    _add_paragraphs(doc, scene["overview"])
                for j, measure in enumerate(scene.get("measures", [])):
                    prefix = ordinals[j] if j < len(ordinals) else f"第{j+1}，"
                    p = doc.add_paragraph()
                    _set_paragraph_format(p)
                    run = p.add_run(f"{prefix}，{_strip_inline_sources(measure)}")
                    _set_run_font(run, FONT_FANGSONG, 16)
        else:
            p = doc.add_paragraph()
            _set_paragraph_format(p)
            run = p.add_run("【待补充场景建设成效】")
            _set_run_font(run, FONT_FANGSONG, 16)

        _heading(doc, "三、下一步工作计划", 1)
        ordinals = ["一是", "二是", "三是", "四是", "五是", "六是"]
        for i, step in enumerate(s.get("next_steps", [])):
            prefix = ordinals[i] if i < len(ordinals) else f"第{i+1}，"
            p = doc.add_paragraph()
            _set_paragraph_format(p)
            run = p.add_run(f"{prefix}，{_strip_inline_sources(step)}")
            _set_run_font(run, FONT_FANGSONG, 16)

        if s.get("conclusion"):
            _add_paragraphs(doc, s["conclusion"])

        _add_common_appendices(doc, customer_type, region, s)
        return _save_doc(doc, output_path, region, project_name)

    # ───── 一、现状分析 ─────
    _heading(doc, "一、现状分析", 1)

    _heading(doc, _subheading(1, "政策背景"), 2)
    _add_paragraphs(doc, s.get("policy_background", "【待补充】"))

    _heading(doc, _subheading(2, "现状概述"), 2)
    _add_paragraphs(doc, s.get("current_status", "【待调研补充】"))

    _heading(doc, _subheading(3, "痛点问题"), 2)
    ordinals = ["一是", "二是", "三是", "四是", "五是", "六是", "七是"]
    for i, pt in enumerate(s.get("pain_points", [])):
        prefix = ordinals[i] if i < len(ordinals) else f"第{i+1}，"
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        run = p.add_run(f"{prefix}，{pt}")
        _set_run_font(run, FONT_FANGSONG, 16)

    # ───── 二、建设目标 ─────
    _heading(doc, "二、建设目标", 1)

    _heading(doc, _subheading(1, "总体目标"), 2)
    _add_paragraphs(doc, s.get("overall_goal", "【待补充】"))

    _heading(doc, _subheading(2, "分项目标"), 2)
    for g in s.get("sub_goals", []):
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        run = p.add_run(g)
        _set_run_font(run, FONT_FANGSONG, 16)

    # ───── 三、建设内容 ─────
    _heading(doc, "三、建设内容", 1)
    for i, mod in enumerate(s.get("modules", []), 1):
        _heading(doc, _subheading(i, mod["name"]), 2)
        _add_paragraphs(doc, mod.get("content", ""))
        prod_p = doc.add_paragraph()
        _set_paragraph_format(prod_p)
        run = prod_p.add_run(f"【产品匹配】{mod.get('product', '待确认')}")
        _set_run_font(run, FONT_FANGSONG, 16, True)

    # ───── 四、技术架构 ─────
    _heading(doc, "四、技术架构", 1)

    _heading(doc, _subheading(1, "总体架构"), 2)
    arch = s.get("architecture", "")
    if arch:
        _add_paragraphs(doc, arch)
    else:
        for layer in [
            "基础设施层：政务云、网络、计算存储等基础资源",
            "数据资源层：数据采集、治理、共享交换、数据中台",
            "应用支撑层：统一身份认证、工作流引擎、消息中心、开发框架",
            "业务应用层：各业务子系统与应用模块",
            "展示层：领导驾驶舱、门户网站、移动端",
        ]:
            p = doc.add_paragraph()
            _set_paragraph_format(p)
            run = p.add_run(layer)
            _set_run_font(run, FONT_FANGSONG, 16)
        doc.add_paragraph("")
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        run = p.add_run("两大支撑体系：标准规范体系、安全保障体系。")
        _set_run_font(run, FONT_FANGSONG, 16)

    _heading(doc, _subheading(2, "技术选型"), 2)
    _add_paragraphs(doc, s.get("tech_selection", "【待补充技术选型说明】"))

    # ───── 五、实施计划 ─────
    _heading(doc, "五、实施计划", 1)

    _heading(doc, _subheading(1, "分期实施路径"), 2)
    phases = s.get("phases", [])
    if phases:
        t = doc.add_table(rows=len(phases) + 1, cols=4)
        _format_table(t)
        _make_header_row(t, ["阶段", "周期", "建设内容", "交付物"])
        for i, ph in enumerate(phases, 1):
            t.rows[i].cells[0].text = ph.get("name", "")
            t.rows[i].cells[1].text = ph.get("duration", "")
            t.rows[i].cells[2].text = ph.get("content", "")
            t.rows[i].cells[3].text = ph.get("deliverables", "")
        _format_table(t)

    _heading(doc, _subheading(2, "项目组织"), 2)
    for role in ["项目领导小组", "项目管理办公室（PMO）", "技术实施团队", "业务配合团队"]:
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        run = p.add_run(role)
        _set_run_font(run, FONT_FANGSONG, 16)

    # ───── 六、预算框架 ─────
    _heading(doc, "六、预算框架", 1)
    items = s.get("budget_items", [])
    if items:
        t = doc.add_table(rows=len(items) + 2, cols=4)
        _format_table(t)
        _make_header_row(t, ["序号", "建设内容", "预算（万元）", "备注"])
        for i, item in enumerate(items, 1):
            t.rows[i].cells[0].text = str(i)
            t.rows[i].cells[1].text = item.get("name", "")
            t.rows[i].cells[2].text = item.get("amount", "待定")
            t.rows[i].cells[3].text = item.get("note", "")
        last = t.rows[-1]
        last.cells[1].text = "合计"
        last.cells[2].text = "待定"
        for c in last.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    _set_run_font(r, FONT_HEITI, 16, True)
        _format_table(t)

    warn = doc.add_paragraph()
    _set_paragraph_format(warn)
    run = warn.add_run("注：预算为估算框架，具体金额由商务部门确认。")
    _set_run_font(run, FONT_FANGSONG, 16)

    # ───── 七、亮点与预期效益 ─────
    _heading(doc, "七、亮点与预期效益", 1)

    _heading(doc, _subheading(1, "方案亮点"), 2)
    for h in s.get("highlights", []):
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        run = p.add_run(h)
        _set_run_font(run, FONT_FANGSONG, 16)

    _heading(doc, _subheading(2, "预期效益"), 2)
    for b in s.get("benefits", []):
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        run = p.add_run(b)
        _set_run_font(run, FONT_FANGSONG, 16)

    _heading(doc, _subheading(3, "成功案例参考"), 2)
    cases = s.get("cases", [])
    if cases:
        for case in cases:
            p = doc.add_paragraph()
            _set_paragraph_format(p)
            run = p.add_run(case.get("name", ""))
            _set_run_font(run, FONT_FANGSONG, 16, True)
            if case.get("summary"):
                p = doc.add_paragraph()
                _set_paragraph_format(p)
                run = p.add_run(_strip_inline_sources(case["summary"]))
                _set_run_font(run, FONT_FANGSONG, 16)
    else:
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        run = p.add_run("【待补充成功案例】")
        _set_run_font(run, FONT_FANGSONG, 16)

    _add_common_appendices(doc, customer_type, region, s)

    # ───── 保存 ─────
    return _save_doc(doc, output_path, region, project_name)


# ── CLI 入口 ──────────────────────────────────────────────

def main():
    start_time = time.perf_counter()
    parser = argparse.ArgumentParser(description="ZX-01 方案初稿 Word 生成器")
    parser.add_argument("--json", default=None, help="输入 JSON 文件路径")
    parser.add_argument("--output", default=None, help="输出 .docx 路径（可选）")
    parser.add_argument("--schema", action="store_true", help="打印输入 JSON 结构说明")
    args = parser.parse_args()

    if args.schema:
        print(SCHEMA)
        return

    if not args.json:
        parser.error("the following arguments are required: --json")

    with open(args.json, "r", encoding="utf-8") as f:
        data = json.load(f)

    path = build_proposal_docx(data, args.output)
    elapsed = time.perf_counter() - start_time
    print(path)
    print(f"generated_in_seconds={elapsed:.2f}")


if __name__ == "__main__":
    main()

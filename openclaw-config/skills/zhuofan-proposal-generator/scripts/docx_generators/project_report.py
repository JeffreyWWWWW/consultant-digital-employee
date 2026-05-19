import os
import re
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FONT_FANGSONG = "仿宋_GB2312"
FONT_HEITI = "黑体"
FONT_XIAOBIAOSONG = "方正小标宋_GBK"
BODY_FONT_SIZE_PT = 16
FIRST_LINE_INDENT_CHARS = 2
REVIEW_MARKERS = (
    "待核验",
    "待确认",
    "待补充",
    "待测算",
    "需人工审核",
    "需人工复核",
    "推断",
    "建议",
    "未核验",
    "模型生成",
    "模型建议",
)
REVIEW_COMMENT_NOTE = "文中批注内容为需人工核对事项，请结合原始材料、政策原文和客户确认结果复核。"
CN_NUMERALS = ("一", "二", "三", "四", "五", "六", "七", "八", "九", "十")
COMMENTS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
NS_CT = "http://schemas.openxmlformats.org/package/2006/content-types"
INLINE_SOURCE_RE = re.compile(r"[（(\[]\s*来源\s*[:：]\s*([^）)\]]+?)\s*[）)\]]")


def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_format(p, first_line_indent: bool = True, space_before: int = 0, space_after: int = 0):
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if first_line_indent:
        ind = p._p.get_or_add_pPr().get_or_add_ind()
        ind.set(qn("w:firstLineChars"), str(FIRST_LINE_INDENT_CHARS * 100))
        if ind.get(qn("w:firstLine")) is not None:
            del ind.attrib[qn("w:firstLine")]


def _setup_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def _add_run(p, text: str, font: str = FONT_FANGSONG, size: int = BODY_FONT_SIZE_PT, bold: bool = False):
    run = p.add_run(text or "")
    _set_run_font(run, font, size, bold)
    return run


def _add_hyperlink(p, text: str, url: str):
    part = p.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text or url
    run.append(text_element)
    hyperlink.append(run)
    p._p.append(hyperlink)


def _heading(doc, text: str, level: int):
    p = doc.add_paragraph()
    _set_paragraph_format(p, first_line_indent=False, space_before=8, space_after=4)
    if level == 1:
        _add_run(p, text, FONT_HEITI, BODY_FONT_SIZE_PT, True)
    else:
        _add_run(p, text, FONT_FANGSONG, BODY_FONT_SIZE_PT, True)
    return p


def _cn_number(index: int) -> str:
    if 1 <= index <= 10:
        return CN_NUMERALS[index - 1]
    if 11 <= index <= 19:
        return "十" + CN_NUMERALS[index - 11]
    if index == 20:
        return "二十"
    return str(index)


def _numbered_subheading(index: int, text: str) -> str:
    text = (text or "").strip()
    if re.match(r"^（[一二三四五六七八九十百千万]+）", text):
        return text
    if re.match(r"^[一二三四五六七八九十百千万]+、", text):
        return text
    return f"（{_cn_number(index)}）{text}"


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _needs_review_comment(text: str) -> bool:
    return any(marker in (text or "") for marker in REVIEW_MARKERS)


def _comment_store(paragraph):
    part = paragraph.part
    if not hasattr(part, "_zhuofan_comments"):
        part._zhuofan_comments = []
    return part._zhuofan_comments


def _add_comment_to_paragraph(paragraph, comment_text: str):
    comment_text = str(comment_text or "").strip()
    if not comment_text:
        return
    if not paragraph.runs:
        _add_run(paragraph, "")
    comments = _comment_store(paragraph)
    comment_id = len(comments)
    comments.append({"id": comment_id, "text": comment_text})

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))
    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)

    p = paragraph._p
    insert_at = 1 if p.pPr is not None else 0
    p.insert(insert_at, start)
    p.append(end)
    p.append(ref_run)


def _review_comment_text(text: str) -> str:
    markers = [marker for marker in REVIEW_MARKERS if marker in (text or "")]
    if markers:
        return "需人工核对：文本包含审核标记（" + "、".join(markers) + "）。"
    return "需人工核对：该内容需结合原始材料和来源依据复核。"


def _add_text_run(p, text: str, review_comment: bool = True):
    return _add_run(p, text)


def _add_paragraphs(doc, text, review_comment: bool = True):
    if not text:
        return
    if isinstance(text, list):
        for item in text:
            _add_paragraphs(doc, item, review_comment=review_comment)
        return
    for para in str(text).split("\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        _add_text_run(p, para, review_comment=review_comment)


def _comment_terms_in_document(doc, terms):
    terms = [str(term).strip() for term in (terms or []) if str(term).strip()]
    if not terms:
        return
    for paragraph in doc.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text or not any(term in full_text for term in terms):
            continue
        matched_terms = [term for term in terms if term in full_text]
        _add_comment_to_paragraph(paragraph, "需人工核对：" + "；".join(matched_terms))


def _paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def _replace_paragraph_text(paragraph, text: str):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    _add_run(paragraph, text)


def _host_from_url(url: str) -> str:
    match = re.search(r"https?://([^/\s]+)", url or "")
    return match.group(1).lower() if match else ""


def _clean_title(project_name: str) -> str:
    name = (project_name or "").strip(" 　《》")
    name = re.sub(r"(解决方案|方案初稿|初稿)$", "", name).strip(" 　-_")
    return name


def _project_report_title(project_name: str, customer_name: str = "") -> str:
    name = _clean_title(project_name)
    customer = (customer_name or "").strip(" 　《》")
    if "汇报" in name:
        if customer and customer not in name:
            return f"{customer}{name}"
        return name
    if name.endswith("建设项目"):
        report_name = f"{name}汇报"
    elif name.endswith("建设"):
        report_name = f"{name}项目汇报"
    else:
        report_name = f"{name}建设项目汇报"
    if customer and customer not in report_name:
        return f"{customer}{report_name}"
    return report_name


def _safe_filename(text: str) -> str:
    text = (text or "").replace("/", "_").replace("\\", "_").replace(" ", "")
    return re.sub(r'[<>:"|?*]', "_", text)


def _default_output_path(region: str, project_name: str, customer_name: str = "") -> str:
    today = date.today().strftime("%Y%m%d")
    report_name = _safe_filename(_project_report_title(project_name, customer_name))
    return f"{report_name}_{today}.docx"


def _normalize_output_path(output_path: str, region: str, project_name: str, customer_name: str = "") -> str:
    filename = _default_output_path(region, project_name, customer_name)
    if not output_path:
        return filename
    output_dir = os.path.dirname(output_path) or "."
    return os.path.join(output_dir, filename)


def _project_contents(sections: dict):
    return sections.get("project_contents") or sections.get("modules") or []


def _format_source_item(item) -> str:
    if not isinstance(item, dict):
        return str(item)
    name = item.get("name") or item.get("title") or item.get("source_name") or item.get("policy_name") or "未命名来源"
    agency = item.get("agency") or item.get("publisher") or item.get("source") or item.get("source_org")
    doc_no = item.get("doc_no") or item.get("document_no")
    date_text = item.get("date") or item.get("publish_date") or item.get("published_at")
    url = item.get("url") or item.get("link") or item.get("source_url")
    status = item.get("status") or item.get("verification_status") or item.get("核验状态")
    used_for = item.get("used_for") or item.get("support") or item.get("purpose")
    parts = [str(name)]
    meta = [value for value in (agency, doc_no, date_text) if value]
    if meta:
        parts.append("（" + "，".join(map(str, meta)) + "）")
    if url:
        parts.append(f"：{url}")
    elif status:
        parts.append(f"：{status}")
    else:
        parts.append("：待核验原文")
    if used_for:
        parts.append(f"；支撑内容：{used_for}")
    return "".join(parts)


def _source_comment_text(item) -> str:
    if not isinstance(item, dict):
        return str(item)
    name = item.get("name") or item.get("title") or item.get("source_name") or item.get("policy_name") or "未命名来源"
    agency = item.get("agency") or item.get("publisher") or item.get("source") or item.get("source_org")
    doc_no = item.get("doc_no") or item.get("document_no")
    date_text = item.get("date") or item.get("publish_date") or item.get("published_at")
    url = item.get("url") or item.get("link") or item.get("source_url")
    used_for = item.get("used_for") or item.get("support") or item.get("purpose")
    meta = "，".join(str(value) for value in (agency, doc_no, date_text) if value)
    title = f"{name}（{meta}）" if meta else str(name)
    parts = [title]
    if url:
        parts.append(str(url))
    if used_for:
        parts.append(f"支撑内容：{used_for}")
    return "；".join(parts)


def _source_fields(item):
    if not isinstance(item, dict):
        return {"text": str(item), "url": ""}
    source_names = item.get("source_names") or []
    source_urls = item.get("source_urls") or []
    if isinstance(source_names, str):
        source_names = [source_names]
    if isinstance(source_urls, str):
        source_urls = [source_urls]
    name = (
        item.get("name")
        or item.get("title")
        or item.get("source_name")
        or item.get("policy_name")
        or (source_names[0] if source_names else "")
        or "未命名来源"
    )
    agency = item.get("agency") or item.get("publisher") or item.get("source") or item.get("source_org")
    doc_no = item.get("doc_no") or item.get("document_no")
    date_text = item.get("date") or item.get("publish_date") or item.get("published_at")
    url = item.get("url") or item.get("link") or item.get("source_url") or (source_urls[0] if source_urls else "")
    status = item.get("status") or item.get("verification_status") or item.get("核验状态")
    used_for = item.get("used_for") or item.get("support") or item.get("purpose")
    meta = [value for value in (agency, doc_no, date_text) if value]
    prefix = f"- {name}"
    if meta:
        prefix += "（" + "，".join(map(str, meta)) + "）"
    suffix = ""
    if used_for:
        suffix += f"；支撑内容：{used_for}"
    if not url:
        suffix = f"：{status or '待核验原文'}" + suffix
    return {"prefix": prefix, "url": url or "", "suffix": suffix}


def _source_items(items):
    if not items:
        return []
    if isinstance(items, str):
        return [line.strip() for line in items.split("\n") if line.strip()]
    return list(items)


def _add_bullets(doc, items):
    if not items:
        return
    if isinstance(items, str):
        _add_paragraphs(doc, items)
        return
    for item in items:
        text = _format_source_item(item)
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        _add_run(p, f"- {text}")


def _add_numbered_review_items(doc, items):
    if not items:
        return
    if isinstance(items, str):
        items = [item.strip() for item in items.split("\n") if item.strip()]
    for idx, item in enumerate(items, start=1):
        text = _review_note_text(item)
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        _add_text_run(p, f"{idx}. {text}", review_comment=False)


def _add_source_bullets(doc, items):
    if not items:
        return
    if isinstance(items, str):
        _add_paragraphs(doc, items)
        return
    for idx, item in enumerate(items, start=1):
        fields = _source_fields(item)
        if "text" in fields:
            p = doc.add_paragraph()
            _set_paragraph_format(p)
            _add_run(p, fields["text"])
            continue
        title = re.sub(r"^\s*-\s*", "", fields["prefix"])
        p = doc.add_paragraph()
        _set_paragraph_format(p)
        _add_run(p, f"{idx}. {title}")

        p = doc.add_paragraph()
        _set_paragraph_format(p)
        _add_run(p, "原文链接：")
        if fields["url"]:
            _add_hyperlink(p, fields["url"], fields["url"])
            _add_run(p, "；")
        else:
            status = fields["suffix"].lstrip("：").split("；", 1)[0] or "待核验原文"
            _add_run(p, f"{status}；")

        support = fields["suffix"].split("支撑内容：", 1)[1] if "支撑内容：" in fields["suffix"] else ""
        if support:
            p = doc.add_paragraph()
            _set_paragraph_format(p)
            _add_run(p, f"支撑内容：{support}")


def _comment_inline_sources_in_body(doc, sources):
    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        matches = INLINE_SOURCE_RE.findall(text)
        if not matches:
            continue
        clean_text = INLINE_SOURCE_RE.sub("", text)
        clean_text = re.sub(r"\s+([，。；：])", r"\1", clean_text).strip()
        _replace_paragraph_text(paragraph, clean_text)


def _review_note_text(item) -> str:
    if isinstance(item, dict):
        text = item.get("content") or item.get("text") or item.get("note") or item.get("issue") or _format_source_item(item)
    else:
        text = str(item)
    text = re.sub(r"^\s*[-•]\s*", "", text)
    text = re.sub(r"^\s*\d+[\.、]\s*", "", text)
    return text.strip()


def _review_note_targets(item):
    if isinstance(item, dict):
        raw_targets = item.get("target") or item.get("targets") or item.get("match_text") or item.get("body_text")
        if isinstance(raw_targets, str):
            return [raw_targets.strip()] if raw_targets.strip() else []
        if isinstance(raw_targets, list):
            return [str(target).strip() for target in raw_targets if str(target).strip()]

    text = _review_note_text(item)
    targets = []
    targets.extend(re.findall(r"[\u4e00-\u9fff]+〔\d{4}〕\d+号", text))
    targets.extend(re.findall(r"[《“\"]([^》”\"]{2,60})[》”\"]", text))
    targets.extend(re.findall(r"[A-Za-z0-9一-龥]+(?:\d+%|％)", text))
    prefix = re.split(r"(?:需|待|建议|为|涉及|由|，|。)", text, maxsplit=1)[0].strip()
    if 2 <= len(prefix) <= 40:
        targets.append(prefix)
    for token in re.findall(r"[A-Za-z0-9\u4e00-\u9fff]{2,30}", text):
        if len(token) >= 4 and token not in ("人工审核事项", "原文链接", "会议纪要", "进一步确认"):
            targets.append(token)
    deduped = []
    for target in targets:
        target = str(target).strip(" ：:，。；;")
        if target and target not in deduped:
            deduped.append(target)
    return deduped


def _source_match_text(item) -> str:
    if not isinstance(item, dict):
        return str(item)
    values = []
    for key in (
        "name",
        "title",
        "source_name",
        "policy_name",
        "agency",
        "publisher",
        "source",
        "source_org",
        "doc_no",
        "document_no",
        "date",
        "publish_date",
        "url",
        "link",
        "source_url",
    ):
        value = item.get(key)
        if value:
            values.append(str(value))
    return " ".join(values)


def _matching_source_for_review(note: str, targets, sources):
    source_items = [item for item in _source_items(sources) if isinstance(item, dict)]
    tokens = [note] + [target for target in targets if target]
    for item in source_items:
        source_text = _source_match_text(item)
        for token in tokens:
            token = str(token or "").strip()
            if len(token) >= 4 and token in source_text:
                return item
    return None


def _source_suffix_for_review(note: str, targets, sources) -> str:
    source = _matching_source_for_review(note, targets, sources)
    if not source:
        return ""
    url = source.get("url") or source.get("link") or source.get("source_url")
    if not url:
        return ""
    return "；可核验来源：" + _source_comment_text(source)


def _review_item_count(review_notes) -> int:
    return len(_source_items(review_notes))


def _comment_review_notes_in_body(doc, review_notes, sources=None):
    unmatched = []
    comment_count = 0
    for idx, item in enumerate(_source_items(review_notes), start=1):
        note = _review_note_text(item)
        if not note:
            continue
        targets = _review_note_targets(item)
        comment_text = f"需人工核对（附录B第{idx}项）：{note}{_source_suffix_for_review(note, targets, sources)}"
        matched = False
        for paragraph in doc.paragraphs:
            text = _paragraph_text(paragraph)
            if not text:
                continue
            if any(target and target in text for target in targets):
                _add_comment_to_paragraph(paragraph, comment_text)
                comment_count += 1
                matched = True
                break
        if not matched:
            unmatched.append(note)
    if unmatched:
        details = "；".join(unmatched)
        raise ValueError(
            "附录B人工审核事项必须对应前文正文批注，请为以下事项补充正文内容或在 review_notes 中提供 target/match_text："
            + details
        )
    return comment_count


def _append_review_appendix(doc, sections: dict):
    sources = (
        sections.get("policy_sources")
        or sections.get("source_links")
        or sections.get("sources")
        or sections.get("source_note")
    )
    review_notes = sections.get("review_notes") or []

    _page_break(doc)
    _heading(doc, "附录A：政策与资料来源链接", 1)
    if sources:
        _add_source_bullets(doc, sources)
    else:
        _add_paragraphs(doc, "【待补充政策与资料来源链接】")

    _page_break(doc)
    _heading(doc, "附录B：需人工审核事项", 1)
    _add_paragraphs(doc, REVIEW_COMMENT_NOTE, review_comment=False)
    if review_notes:
        _add_numbered_review_items(doc, review_notes)
    else:
        _add_paragraphs(doc, "【待补充需人工审核事项】")


def _validate_source_appendix_items(sources):
    if not sources or isinstance(sources, str):
        return
    issues = []
    for idx, item in enumerate(_source_items(sources), start=1):
        if not isinstance(item, dict):
            continue
        name = item.get("name") or item.get("title") or item.get("source_name") or item.get("policy_name")
        url = item.get("url") or item.get("link") or item.get("source_url")
        status = item.get("status") or item.get("verification_status") or item.get("核验状态")
        used_for = item.get("used_for") or item.get("support") or item.get("purpose")
        if not name:
            issues.append(f"附录A第{idx}项缺少政策或资料名称")
        if not url and not status:
            issues.append(f"附录A第{idx}项缺少原文链接或核验状态")
        if not used_for:
            issues.append(f"附录A第{idx}项缺少支撑内容")
    if issues:
        raise ValueError("附录A来源清单不完整：" + "；".join(issues))


def _source_is_verified(item) -> bool:
    if not isinstance(item, dict):
        return False
    status = str(item.get("status") or item.get("verification_status") or item.get("核验状态") or "").strip()
    return status in {"已核验原文", "人工已核验原文", "已核验官网原文"}


def _source_review_tokens(item):
    if not isinstance(item, dict):
        return [str(item).strip()]
    tokens = []
    for key in ("name", "title", "source_name", "policy_name", "doc_no", "document_no", "url", "link", "source_url"):
        value = item.get(key)
        if value:
            tokens.append(str(value).strip())
    return [token for token in tokens if len(token) >= 4]


def _review_note_match_text(item) -> str:
    if isinstance(item, dict):
        values = [
            item.get("content"),
            item.get("text"),
            item.get("note"),
            item.get("issue"),
            item.get("target"),
            item.get("match_text"),
            item.get("body_text"),
        ]
        targets = item.get("targets")
        if isinstance(targets, list):
            values.extend(targets)
        return " ".join(str(value) for value in values if value)
    return str(item)


def _validate_source_review_coverage(sources, review_notes):
    if not sources or isinstance(sources, str):
        return
    review_text = "\n".join(_review_note_match_text(item) for item in _source_items(review_notes))
    missing = []
    for idx, item in enumerate(_source_items(sources), start=1):
        if not isinstance(item, dict) or _source_is_verified(item):
            continue
        tokens = _source_review_tokens(item)
        if not tokens or not any(token in review_text for token in tokens):
            name = item.get("name") or item.get("title") or item.get("source_name") or item.get("policy_name") or f"第{idx}项来源"
            missing.append(f"附录A第{idx}项 {name}")
    if missing:
        raise ValueError(
            "附录A来源必须逐条审查：以下来源未标记为“已核验原文”，也未进入附录B人工审核事项："
            + "；".join(map(str, missing))
        )


def _comments_xml(comments):
    ET.register_namespace("w", NS_W)
    root = ET.Element(f"{{{NS_W}}}comments")
    for item in comments:
        comment = ET.SubElement(
            root,
            f"{{{NS_W}}}comment",
            {
                f"{{{NS_W}}}id": str(item["id"]),
                f"{{{NS_W}}}author": "卓智",
                f"{{{NS_W}}}initials": "ZZ",
            },
        )
        p = ET.SubElement(comment, f"{{{NS_W}}}p")
        r = ET.SubElement(p, f"{{{NS_W}}}r")
        t = ET.SubElement(r, f"{{{NS_W}}}t")
        t.text = item["text"]
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _ensure_comments_relationship(xml_bytes):
    ET.register_namespace("", NS_REL)
    root = ET.fromstring(xml_bytes)
    for rel in root.findall(f"{{{NS_REL}}}Relationship"):
        if rel.get("Type") == COMMENTS_REL_TYPE:
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)
    existing_ids = {rel.get("Id") for rel in root.findall(f"{{{NS_REL}}}Relationship")}
    idx = 1
    while f"rId{idx}" in existing_ids:
        idx += 1
    ET.SubElement(
        root,
        f"{{{NS_REL}}}Relationship",
        {"Id": f"rId{idx}", "Type": COMMENTS_REL_TYPE, "Target": "comments.xml"},
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _ensure_comments_content_type(xml_bytes):
    ET.register_namespace("", NS_CT)
    root = ET.fromstring(xml_bytes)
    for override in root.findall(f"{{{NS_CT}}}Override"):
        if override.get("PartName") == "/word/comments.xml":
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)
    ET.SubElement(
        root,
        f"{{{NS_CT}}}Override",
        {"PartName": "/word/comments.xml", "ContentType": COMMENTS_CONTENT_TYPE},
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _write_comments_part(docx_path: str, comments):
    if not comments:
        return
    with zipfile.ZipFile(docx_path, "r") as src:
        entries = {name: src.read(name) for name in src.namelist()}

    entries["word/comments.xml"] = _comments_xml(comments)
    entries["word/_rels/document.xml.rels"] = _ensure_comments_relationship(entries["word/_rels/document.xml.rels"])
    entries["[Content_Types].xml"] = _ensure_comments_content_type(entries["[Content_Types].xml"])

    output_dir = os.path.dirname(os.path.abspath(docx_path)) or "."
    fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=output_dir)
    os.close(fd)
    try:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
            for name, content in entries.items():
                dst.writestr(name, content)
        os.replace(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def build_docx(data: dict, output_path: str = None) -> str:
    project_name = data["project_name"]
    customer_name = data.get("customer_name", "")
    region = data.get("region", "")
    sections = data.get("sections", {})
    sources = (
        sections.get("policy_sources")
        or sections.get("source_links")
        or sections.get("sources")
        or sections.get("source_note")
    )
    review_notes = sections.get("review_notes") or []

    doc = Document()
    _setup_page(doc)

    style = doc.styles["Normal"]
    style.font.name = FONT_FANGSONG
    style.font.size = Pt(BODY_FONT_SIZE_PT)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FANGSONG)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(title_p, first_line_indent=False, space_after=18)
    _add_run(title_p, _project_report_title(project_name, customer_name), FONT_XIAOBIAOSONG, 22, True)

    _heading(doc, "一、项目建设的依据", 1)
    _add_paragraphs(doc, sections.get("project_basis") or sections.get("policy_background") or "【待补充项目建设依据】")

    _heading(doc, "二、建设目标", 1)
    _add_paragraphs(doc, sections.get("project_goals") or sections.get("overall_goal") or "【待补充建设目标】")

    _heading(doc, "三、建设内容", 1)
    contents = _project_contents(sections)
    if isinstance(contents, list) and contents:
        for idx, item in enumerate(contents, start=1):
            if isinstance(item, dict):
                name = item.get("name") or item.get("title")
                content = item.get("content") or item.get("description") or item.get("summary")
                if name:
                    _heading(doc, _numbered_subheading(idx, name), 2)
                _add_paragraphs(doc, content)
            else:
                _add_paragraphs(doc, item)
    else:
        _add_paragraphs(doc, contents or "【待补充建设内容】")

    _validate_source_appendix_items(sources)
    _validate_source_review_coverage(sources, review_notes)
    _comment_inline_sources_in_body(doc, sources)
    review_comment_count = _comment_review_notes_in_body(doc, review_notes, sources)
    review_item_count = _review_item_count(review_notes)
    total_comment_count = len(getattr(doc.part, "_zhuofan_comments", []))
    if total_comment_count != review_item_count or review_comment_count != review_item_count:
        raise ValueError(
            f"附录B人工审核事项与正文批注必须一对一：附录B {review_item_count} 条，"
            f"正文批注 {total_comment_count} 条。请检查 review_notes、target/match_text 和自动批注来源。"
        )

    _append_review_appendix(doc, sections)

    output_path = _normalize_output_path(output_path, region, project_name, customer_name)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    _write_comments_part(output_path, getattr(doc.part, "_zhuofan_comments", []))
    return os.path.abspath(output_path)

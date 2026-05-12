"""
ZX-01 方案初稿 Word 文档生成器

将结构化的方案内容生成为符合公司标准的 .docx 文件。
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date
import os


def set_cell_shading(cell, color_hex: str):
    """设置表格单元格背景色"""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    tc_pr.append(shading_elm)


def add_heading_styled(doc: Document, text: str, level: int):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)  # 深蓝
    return heading


def build_proposal(
    project_name: str,
    customer_name: str,
    customer_type: str,
    region: str,
    sections: dict,
    output_path: str = None,
) -> str:
    """
    生成方案初稿 Word 文档。

    Args:
        project_name: 项目名称，如"城市大脑建设"
        customer_name: 客户单位名称，如"杭州市数据局"
        customer_type: 客户类型 (数据局/政务中心/大数据中心/企业)
        region: 区域，如"杭州市"
        sections: 各章节内容，结构见下方说明
        output_path: 输出路径，默认自动生成

    sections 结构:
    {
        "policy_background": "政策背景文本",
        "current_status": "现状概述文本",
        "pain_points": ["痛点1", "痛点2", ...],
        "overall_goal": "总体目标文本",
        "sub_goals": ["子目标1", "子目标2", ...],
        "modules": [
            {"name": "模块名", "content": "详细内容", "product": "匹配产品"},
            ...
        ],
        "architecture": "技术架构描述",
        "tech_selection": "技术选型说明",
        "phases": [
            {"name": "一期", "duration": "X个月", "deliverables": "交付物"},
            ...
        ],
        "budget_items": [
            {"name": "项目名", "amount": "待定", "note": ""},
            ...
        ],
        "highlights": ["亮点1", "亮点2", ...],
        "benefits": ["效益1", "效益2", ...],
        "cases": [{"name": "案例名", "source": "来源", "summary": "摘要"}],
        "ref_proposals": "参考的历史方案",
        "ref_products": "匹配的产品模块",
        "ref_policies_count": 0,
        "ref_cases_count": 0,
    }

    Returns:
        生成的 .docx 文件路径
    """
    doc = Document()

    # ========== 全局样式 ==========
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.paragraph_format.line_spacing = 1.5

    # ========== 封面信息 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(6)
    run = title.add_run(f"《{project_name}》解决方案")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(24)
    run = subtitle.add_run("（初稿）")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 元信息表格
    meta_table = doc.add_table(rows=4, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("编制单位", "卓繁信息集团", "目标客户", customer_name),
        ("客户类型", customer_type, "区域", region),
        ("编制日期", date.today().strftime("%Y-%m-%d"), "版本", "V0.1（AI 初稿）"),
        ("状态", "待人工审核", "", ""),
    ]
    for i, (k1, v1, k2, v2) in enumerate(meta_data):
        cells = meta_table.rows[i].cells
        cells[0].text = k1
        cells[1].text = v1
        if k2:
            cells[2].text = k2
            cells[3].text = v2
        for j in [0, 2]:
            if cells[j].text:
                set_cell_shading(cells[j], "E8EDF5")
                for p in cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.add_paragraph("")  # 空行

    # ========== 一、现状分析 ==========
    add_heading_styled(doc, "一、现状分析", level=1)

    add_heading_styled(doc, "1.1 政策背景", level=2)
    doc.add_paragraph(sections.get("policy_background", "【待补充】"))

    add_heading_styled(doc, "1.2 现状概述", level=2)
    doc.add_paragraph(sections.get("current_status", "【待调研补充】"))

    add_heading_styled(doc, "1.3 痛点问题", level=2)
    pain_points = sections.get("pain_points", [])
    ordinals = ["一是", "二是", "三是", "四是", "五是"]
    for i, point in enumerate(pain_points):
        prefix = ordinals[i] if i < len(ordinals) else f"第{i+1}，"
        doc.add_paragraph(f"{prefix}，{point}", style="List Bullet")

    # ========== 二、建设目标 ==========
    add_heading_styled(doc, "二、建设目标", level=1)

    add_heading_styled(doc, "2.1 总体目标", level=2)
    doc.add_paragraph(sections.get("overall_goal", "【待补充】"))

    add_heading_styled(doc, "2.2 分项目标", level=2)
    for goal in sections.get("sub_goals", []):
        doc.add_paragraph(goal, style="List Bullet")

    # ========== 三、建设内容 ==========
    add_heading_styled(doc, "三、建设内容", level=1)
    for i, module in enumerate(sections.get("modules", []), 1):
        add_heading_styled(doc, f"3.{i} {module['name']}", level=2)
        doc.add_paragraph(module["content"])
        product_info = doc.add_paragraph()
        run = product_info.add_run(f"【产品匹配】{module.get('product', '待匹配')}")
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        run.bold = True

    # ========== 四、技术架构 ==========
    add_heading_styled(doc, "四、技术架构", level=1)

    add_heading_styled(doc, "4.1 总体架构", level=2)
    arch_text = sections.get("architecture", "")
    if arch_text:
        doc.add_paragraph(arch_text)
    else:
        layers = [
            "基础设施层：提供计算、存储、网络等基础资源支撑",
            "数据资源层：实现数据采集、治理、共享与开放",
            "平台支撑层：提供统一的中间件、开发框架与服务能力",
            "业务应用层：承载各类业务应用系统",
            "用户交互层：面向各类用户的统一入口",
        ]
        for layer in layers:
            doc.add_paragraph(layer, style="List Bullet")
        doc.add_paragraph("")
        doc.add_paragraph("两大支撑体系：标准规范体系、安全保障体系。")

    add_heading_styled(doc, "4.2 技术选型", level=2)
    doc.add_paragraph(sections.get("tech_selection", "【待补充技术选型说明】"))

    # ========== 五、实施计划 ==========
    add_heading_styled(doc, "五、实施计划", level=1)

    add_heading_styled(doc, "5.1 实施路径", level=2)
    phases = sections.get("phases", [])
    if phases:
        phase_table = doc.add_table(rows=len(phases) + 1, cols=3)
        phase_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["阶段", "周期", "交付物"]
        for j, h in enumerate(headers):
            phase_table.rows[0].cells[j].text = h
            set_cell_shading(phase_table.rows[0].cells[j], "1A3C6E")
            for p in phase_table.rows[0].cells[j].paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
        for i, phase in enumerate(phases, 1):
            phase_table.rows[i].cells[0].text = phase["name"]
            phase_table.rows[i].cells[1].text = phase["duration"]
            phase_table.rows[i].cells[2].text = phase["deliverables"]

    # ========== 六、预算框架 ==========
    add_heading_styled(doc, "六、预算框架", level=1)
    budget_items = sections.get("budget_items", [])
    if budget_items:
        budget_table = doc.add_table(rows=len(budget_items) + 2, cols=4)
        budget_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        b_headers = ["序号", "建设内容", "预算（万元）", "备注"]
        for j, h in enumerate(b_headers):
            budget_table.rows[0].cells[j].text = h
            set_cell_shading(budget_table.rows[0].cells[j], "1A3C6E")
            for p in budget_table.rows[0].cells[j].paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
        for i, item in enumerate(budget_items, 1):
            budget_table.rows[i].cells[0].text = str(i)
            budget_table.rows[i].cells[1].text = item["name"]
            budget_table.rows[i].cells[2].text = item.get("amount", "待定")
            budget_table.rows[i].cells[3].text = item.get("note", "")
        # 合计行
        last = budget_table.rows[-1]
        last.cells[0].text = ""
        last.cells[1].text = "合计"
        last.cells[2].text = "待定"
        last.cells[3].text = ""
        for cell in last.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

    warn = doc.add_paragraph()
    run = warn.add_run("⚠️ 预算为估算框架，具体金额需人工确认。")
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

    # ========== 七、亮点与预期效益 ==========
    add_heading_styled(doc, "七、亮点与预期效益", level=1)

    add_heading_styled(doc, "7.1 方案亮点", level=2)
    for h in sections.get("highlights", []):
        doc.add_paragraph(h, style="List Bullet")

    add_heading_styled(doc, "7.2 预期效益", level=2)
    for b in sections.get("benefits", []):
        doc.add_paragraph(b, style="List Bullet")

    add_heading_styled(doc, "7.3 成功案例参考", level=2)
    cases = sections.get("cases", [])
    if cases:
        for case in cases:
            p = doc.add_paragraph()
            run = p.add_run(f"■ {case['name']}")
            run.bold = True
            doc.add_paragraph(case.get("summary", ""))
            source_p = doc.add_paragraph()
            run = source_p.add_run(f"【来源】{case.get('source', '知识库')}")
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(9)
    else:
        doc.add_paragraph("【待补充成功案例】")

    # ========== 人工审核提示 ==========
    doc.add_page_break()
    add_heading_styled(doc, "⚠️ 人工审核提示", level=1)
    checklist = [
        "政策引用是否准确（核对原文）",
        "产品模块匹配是否与公司最新产品清单一致",
        "预算估算是否合理",
        "方案创新点是否有差异化",
        "客户特殊需求是否完整覆盖",
    ]
    for item in checklist:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")

    # ========== 生成说明 ==========
    add_heading_styled(doc, "📋 生成说明", level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("参考历史方案", sections.get("ref_proposals", "无")),
        ("匹配产品模块", sections.get("ref_products", "无")),
        ("引用政策数量", f"{sections.get('ref_policies_count', 0)} 条"),
        ("引用案例数量", f"{sections.get('ref_cases_count', 0)} 个"),
        ("生成工具", "卓智 - 咨询顾问数字员工"),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        set_cell_shading(info_table.rows[i].cells[0], "E8EDF5")
        for p in info_table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    # ========== 保存 ==========
    # 文件名规范：ZX01_{区域}_{项目简称}_解决方案_初稿_{日期}.docx
    # 示例：ZX01_杭州市_城市大脑_解决方案_初稿_20260512.docx
    if not output_path:
        today = date.today().strftime("%Y%m%d")
        safe_region = region.replace("/", "_").replace("\\", "_")
        # 取项目名前 10 个字符作为简称
        short_name = project_name[:10].replace("/", "_").replace("\\", "_")
        output_path = f"ZX01_{safe_region}_{short_name}_解决方案_初稿_{today}.docx"

    doc.save(output_path)
    return os.path.abspath(output_path)



# ========== 便捷入口 ==========
if __name__ == '__main__':
    test_sections = {
        'policy_background': (
            '当前，党中央、国务院高度重视数字政府建设。'
            '《国务院关于加强数字政府建设的指导意见》（国发〔２０２２〕14号）明确提出，'
            '要统筹推进各行业各领域政务应用系统集约建设、互联互通、协同联动，'
            '构建协同高效的政府数字化履职能力体系。'
        ),
        'current_status': (
            '杭州市已建成政务云平台和城市大脑基础平台，累计接入各部门应用系统超过200个。'
            '但各部门数据仍存在壁垒，跨部门协同效率有待提升。'
        ),
        'pain_points': [
            '各部门业务系统独立建设，数据标准不统一，形成信息孤岛',
            '基层工作人员需在多个系统重复录入相同数据',
            '城市治理缺乏统一指挥调度平台，应急响应速度不足',
            '数据安全防护体系尚不健全',
        ],
        'overall_goal': (
            '构建杭州市城市大脑综合治理平台，打通各部门数据壁垒，'
            '实现跨部门、跨层级的数据共享与业务协同，全面提升城市治理效能。'
        ),
        'sub_goals': [
            '建成统一的数据共享交换平台，实现核心部门数据100%接入',
            '构建城市运行管理中心',
            '推进基层减负，减少重复填报事飡50%以上',
            '建立健全数据安全防护体系',
        ],
        'modules': [
            {
                'name': '数据共享交换平台',
                'content': '建设统一的数据共享交换平台，实现各部门数据的统一汇聚、标准化治理与按需共享。',
                'product': '卓繁数据共享交换平台',
            },
            {
                'name': '城市运行管理中心',
                'content': '搭建城市运行管理中心，汇聚城管、交通、环保、应急等领域数据。',
                'product': '卓繁城市运管平台',
            },
            {
                'name': '基层减负应用',
                'content': '整合基层高频填报事项，减少基层重复填报工作量。',
                'product': '卓繁基层治理平台',
            },
        ],
        'architecture': '',
        'tech_selection': '全面采用信创技术路线，采用微服务架构，支持弹性扩展与高可用部署。',
        'phases': [
            {'name': '一期', 'duration': '6个月', 'deliverables': '数据共享交换平台上线'},
            {'name': '二期', 'duration': '6个月', 'deliverables': '城市运管中心上线'},
            {'name': '三期', 'duration': '6个月', 'deliverables': '全面推广，运营体系建立'},
        ],
        'budget_items': [
            {'name': '数据共享交换平台', 'amount': '待定'},
            {'name': '城市运行管理中心', 'amount': '待定'},
            {'name': '基层减负应用', 'amount': '待定'},
            {'name': '数据安全体系建设', 'amount': '待定'},
            {'name': '项目管理与培训', 'amount': '待定'},
        ],
        'highlights': [
            '全面对标国家数字政府建设要求',
            '采用信创技术路线，确保系统自主可控',
            '基层减负创新模式',
            '城市运管中心智能化治理',
        ],
        'benefits': [
            '数据共享率从40%提升至90%以上',
            '基层重复填报事项减少50%以上',
            '应急响应时间缩短50%',
            '群众办事满意度提升至95%以上',
        ],
        'cases': [
            {
                'name': '某市城市大脑建设项目',
                'source': '知识库',
                'summary': '接入12个部门数据，实现城市治理效率提升40%。',
            }
        ],
        'ref_proposals': '某市城市大脑建设方案',
        'ref_products': '数据共享交换平台、城市运管平台、基层治理平台',
        'ref_policies_count': 3,
        'ref_cases_count': 1,
    }

    output = build_proposal(
        project_name='杭州市城市大脑综合治理平台建设',
        customer_name='杭州市数据局',
        customer_type='数据局',
        region='杭州市',
        sections=test_sections,
        output_path='D:/repository/consultant-digital-employee/output/ZX01_test_output.docx',
    )
    print(f'方案已生成：{output}')
